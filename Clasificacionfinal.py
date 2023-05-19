import docx
import re
import openpyxl

# Abre el documento de Word
doc = docx.Document('Tema.docx')

# Define la expresión regular para buscar los pasajes de la Biblia
regex = r"(?:(?:\b\w{2,}\s)?\d{1,2}\s?[A-Za-z]{1,3}\b|\b[A-Za-z]{2,}\b)\s+\d+(?:,\d+)?(?:s{1,2}|ss)?(?:\.\d+)?"

# Define una lista de libros permitidos con su respectivo tipo
libros_permitidos = {
    "Gen": "Históricos",
    "Ex": "Históricos",
    "Lv": "Históricos",
    "Nm": "Históricos",
    "Dt": "Históricos",
    "Jos": "Históricos",
    "Jue": "Históricos",
    "Rt": "Históricos",
    "1Sam": "Históricos",
    "2Sam": "Históricos",
    "1Rey": "Históricos",
    "2Rey": "Históricos",
    "1Crón": "Históricos",
    "2Crón": "Históricos",
    "Esdras": "Históricos",
    "Neh": "Históricos",
    "Tob": "Históricos",
    "Jdt": "Históricos",
    "Est": "Históricos",
    "1Mac": "Históricos",
    "2Mac": "Históricos",
     "Is": "Proféticos",
    "Jer": "Proféticos",
    "Lam": "Proféticos",
    "Bar": "Proféticos",
    "Ez": "Proféticos",
    "Dan": "Proféticos",
    "Os": "Proféticos",
    "Jl": "Proféticos",
    "Am": "Proféticos",
    "Abd": "Proféticos",
    "Jon": "Proféticos",
    "Mi": "Proféticos",
    "Nah": "Proféticos",
    "Hab": "Proféticos",
    "Sof": "Proféticos",
    "Ag": "Proféticos",
    "Zac": "Proféticos",
    "Mal": "Proféticos",
    "Mt": "Evangelios",
    "Mc": "Evangelios",
    "Lc": "Evangelios",
    "Jn": "Evangelios",
    "Act": "Cartas",
    "Rom": "Cartas",
    "1Cor": "Cartas",
    "2Cor": "Cartas",
    "Gal": "Cartas",
    "Ef": "Cartas",
    "Flp": "Cartas",
    "Col": "Cartas",
    "1Tes": "Cartas",
    "2Tes": "Cartas",
    "1Tim": "Cartas",
    "2Tim": "Cartas",
    "Tito": "Cartas",
    "Filem": "Cartas",
    "Heb": "Cartas",
    "Sto": "Cartas",
    "1Ped": "Cartas",
    "2Ped": "Cartas",
    "1Jn": "Cartas",
    "2Jn": "Cartas",
    "3Jn": "Cartas",
    "Jud": "Cartas",
    "Sal": "Salmos"
}

# Crea un diccionario para almacenar los pasajes de la Biblia clasificados por tipo
pasajes_clasificados = {
    "Históricos": [],
    "Proféticos": [],
    "Cartas": [],
    "Evangelios": [],
    "Salmos": [],
}

# Recorre todos los párrafos del documento y busca los pasajes de la Biblia
for p in doc.paragraphs:
    matches = re.findall(regex, p.text)
    for match in matches:
        # Elimina los caracteres especiales del pasaje
        pasaje = re.sub(r"\s+", " ", match)
        pasaje = pasaje.replace(".", ",")

        # Verifica si el pasaje es del formato "Libro Capítulo,Versículo-Versículo"
       
       # ...
# Recorre todos los párrafos del documento y busca los pasajes de la Biblia
for p in doc.paragraphs:
    matches = re.findall(regex, p.text)
    for match in matches:
        # Elimina los caracteres especiales del pasaje
        pasaje = re.sub(r"\s+", " ", match)
        pasaje = pasaje.replace(".", ",")

        # Verifica si el pasaje es del formato "Libro Capítulo,Versículo-Versículo"
        if "-" in pasaje:
            partes = pasaje.split(" ")
            libro_capitulo = partes[0]
            versiculos = partes[1].split("-")
            for versiculo in versiculos:
                pasaje_completo = f"{libro_capitulo} {versiculo}"
                if libro_capitulo in libros_permitidos:
                    tipo = libros_permitidos[libro_capitulo]
                    pasajes_clasificados[tipo].append(pasaje_completo)
        elif re.match(r"\b[A-Za-z]{2,}\s\d+,\d+-\d+\b", pasaje):
            partes = pasaje.split(" ")
            libro_capitulo = partes[0]
            versiculos = partes[1].split("-")
            for versiculo in versiculos:
                pasaje_completo = f"{libro_capitulo} {versiculo}"
                if libro_capitulo in libros_permitidos:
                    tipo = libros_permitidos[libro_capitulo]
                    pasajes_clasificados[tipo].append(pasaje_completo)
      # ...

        else:
            # Extrae el libro y el capítulo del pasaje
            partes = pasaje.split(" ")
            libro_capitulo = partes[0]
            
            # Verifica si el libro está permitido
            if libro_capitulo in libros_permitidos:
                tipo = libros_permitidos[libro_capitulo]
                pasajes_clasificados[tipo].append(pasaje)


# Define una función para ordenar alfabéticamente y por capítulo los libros dentro de sus clasificaciones
def ordenar_libros(pasajes_clasificados):
    for tipo, pasajes in pasajes_clasificados.items():
        pasajes.sort(key=lambda x: (x.split()[0], int(x.split()[1].split(",")[0])))  # Ordena los pasajes por libro y capítulo
    return pasajes_clasificados

# ...

# Imprime los pasajes de la Biblia clasificados por tipo (ordenados alfabéticamente y por capítulo)
pasajes_clasificados_ordenados = ordenar_libros(pasajes_clasificados)
for tipo, pasajes in pasajes_clasificados_ordenados.items():
    print(f"--- {tipo} ---")
    for pasaje in pasajes:
        print(pasaje)
    print()



# Imprime los pasajes de la Biblia clasificados por tipo
for tipo, pasajes in pasajes_clasificados.items():
    print(f"--- {tipo} ---")
    for pasaje in pasajes:
        print(pasaje)
    print()

# Crea un nuevo archivo de Excel
workbook = openpyxl.Workbook()
sheet = workbook.active

# Escribe los pasajes de la Biblia clasificados en columnas
column_index = 1
for tipo, pasajes in pasajes_clasificados.items():
    # Escribe el tipo de pasaje en la primera celda de la columna
    sheet.cell(row=1, column=column_index, value=tipo)

    # Escribe los pasajes en las celdas debajo del tipo
    row_index = 2
    for pasaje in pasajes:
        sheet.cell(row=row_index, column=column_index, value=pasaje)
        row_index += 1

    column_index += 1

# Guarda el archivo de Excel con un nombre específico
nombre_archivo = "pasajes_biblicos.xlsx"
workbook.save(nombre_archivo)
print(f"El archivo de Excel '{nombre_archivo}' ha sido creado exitosamente.")